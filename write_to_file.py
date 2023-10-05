import docx
from docx.shared import Cm
from math import ceil


# def write_to_file_with_ans2(mas_var, vol):
#     doc = docx.Document()
#     doc.sections[0].page_height = Cm(29.7)
#     doc.sections[0].page_width = Cm(21.0)
#     doc.sections[0].left_margin = Cm(1.5)
#     doc.sections[0].right_margin = Cm(1.5)
#     doc.sections[0].top_margin = Cm(1.5)
#     doc.sections[0].bottom_margin = Cm(1.5)
#     table = doc.add_table(rows=4 * vol, cols=12)
#
#     for _ in range(0, 4 * vol, 4):
#         table.cell(_, 0).merge(table.cell(_ + 1, 0))
#         table.cell(_, 6).merge(table.cell(_ + 1, 6))
#         table.cell(_ + 2, 1).merge(table.cell(_ + 2, 2))
#         table.cell(_ + 2, 4).merge(table.cell(_ + 2, 5))
#         table.cell(_, 8).merge(table.cell(_, 11))
#         table.cell(_ + 1, 8).merge(table.cell(_ + 1, 11))
#     per = 0
#     for i in range(0, 4 * vol, 4):
#         try:
#             a1, b1, a2, b2, lambda1, lambda2 = mas_var[per]
#         except IndexError:
#             break
#         per += 1
#         table.rows[i].cells[0].paragraphs[0].add_run().add_picture('skoba.png', width=Cm(0.34), height=Cm(1.76))
#         table.rows[i].cells[6].paragraphs[0].add_run().add_picture('skoba.png', width=Cm(0.34), height=Cm(1.76))
#         table.rows[i].cells[1].text = "x`="
#         table.rows[i].cells[2].text = str(a1) + "x"
#         if b1 > 0:
#             table.rows[i].cells[3].text = "+"
#             table.rows[i].cells[4].text = str(b1) + "y"
#         elif b1 < 0:
#             table.rows[i].cells[3].text = "-"
#             table.rows[i].cells[4].text = str(abs(b1)) + "y"
#         else:
#             pass
#         table.rows[i + 1].cells[1].text = "y`="
#         table.rows[i + 1].cells[2].text = str(a2) + "x"
#         if b2 > 0:
#             table.rows[i + 1].cells[3].text = "+"
#             table.rows[i + 1].cells[4].text = str(b2) + "y"
#         elif b2 < 0:
#             table.rows[i + 1].cells[3].text = "-"
#             table.rows[i + 1].cells[4].text = str(abs(b2)) + "y"
#         else:
#             pass
#         table.rows[i + 2].cells[0].text = "\u03bb1="
#         table.rows[i + 2].cells[1].text = str(lambda1)
#         table.rows[i + 2].cells[3].text = "\u03bb2="
#         table.rows[i + 2].cells[4].text = str(lambda2)
#         table.rows[i].cells[10].text = "x(t) ="
#         k11 = 1
#         k21 = (lambda1 - a1) * k11 / b1
#         while k21 != int(k21):
#             k11 += 1
#             k21 = (lambda1 - a1) * k11 / b1
#         k12 = 1
#         k22 = (lambda2 - a1) * k12 / b1
#         while k22 != int(k22):
#             k12 += 1
#             k22 = (lambda2 - a1) * k12 / b1
#         k21 = int(k21)
#         k22 = int(k22)
#         l1 = int(lambda1)
#         l2 = int(lambda2)
#         stroka1 = str(k11) + "C1*exp(" + str(lambda1) + "t) + " + str(k12) + "C2*exp(" + str(lambda2) + "t)"
#         stroka2 = str(k21) + "C1*exp(" + str(lambda1) + "t) + " + str(k22) + "C2*exp(" + str(lambda2) + "t)"
#         table.rows[i].cells[7].text = "x(t) ="
#         table.rows[i + 1].cells[7].text = "y(t) ="
#         table.rows[i].cells[8].text = stroka1
#        table.rows[i + 1].cells[8].text = stroka2
#
#     doc.save('../SDY_2_ANS.docx')


def write_to_file_with_ans2(mas_var, vol):
    doc = docx.Document()
    doc.sections[0].page_height = Cm(29.7)
    doc.sections[0].page_width = Cm(21.0)
    doc.sections[0].left_margin = Cm(1.5)
    doc.sections[0].right_margin = Cm(1.5)
    doc.sections[0].top_margin = Cm(1.5)
    doc.sections[0].bottom_margin = Cm(1.5)
    table = doc.add_table(rows=4 * ceil(vol / 2), cols=15)

    for _ in range(0, 4 * ceil(vol / 2), 4):
        table.cell(_, 0).merge(table.cell(_ + 1, 0))
        table.cell(_, 9).merge(table.cell(_ + 1, 9))
        table.cell(_ + 2, 1).merge(table.cell(_ + 2, 2))
        table.cell(_ + 2, 4).merge(table.cell(_ + 2, 5))
        table.cell(_ + 2, 10).merge(table.cell(_ + 2, 11))
        table.cell(_ + 2, 13).merge(table.cell(_ + 2, 14))
    per = 0
    for i in range(0, 4 * ceil(vol / 2), 4):
        for j in range(0, 13, 9):
            try:
                a1, b1, a2, b2, lambda1, lambda2 = mas_var[per]
            except IndexError:
                break
            per += 1
            table.rows[i].cells[j].paragraphs[0].add_run().add_picture('skoba.png', width=Cm(0.34), height=Cm(1.76))
            table.rows[i].cells[j + 1].text = "x`="
            table.rows[i].cells[j + 2].text = str(a1) + "x"
            if b1 > 0:
                table.rows[i].cells[j + 3].text = "+"
                table.rows[i].cells[j + 4].text = str(b1) + "y"
            elif b1 < 0:
                table.rows[i].cells[j + 3].text = "-"
                table.rows[i].cells[j + 4].text = str(abs(b1)) + "y"
            else:
                pass
            table.rows[i + 1].cells[j + 1].text = "y`="
            table.rows[i + 1].cells[j + 2].text = str(a2) + "x"
            if b2 > 0:
                table.rows[i + 1].cells[j + 3].text = "+"
                table.rows[i + 1].cells[j + 4].text = str(b2) + "y"
            elif b2 < 0:
                table.rows[i + 1].cells[j + 3].text = "-"
                table.rows[i + 1].cells[j + 4].text = str(abs(b2)) + "y"
            else:
                pass
            table.rows[i + 2].cells[j].text = "\u03bb1="
            table.rows[i + 2].cells[j + 1].text = str(lambda1)
            table.rows[i + 2].cells[j + 3].text = "\u03bb2="
            table.rows[i + 2].cells[j + 4].text = str(lambda2)
    doc.save('../Diff2_with_ans.docx')


def write_to_file_without_ans2(mas_var, vol):
    doc = docx.Document()
    doc.sections[0].page_height = Cm(29.7)
    doc.sections[0].page_width = Cm(21.0)
    doc.sections[0].left_margin = Cm(1.5)
    doc.sections[0].right_margin = Cm(1.5)
    doc.sections[0].top_margin = Cm(1.5)
    doc.sections[0].bottom_margin = Cm(1.5)
    table = doc.add_table(rows=3 * ceil(vol / 2), cols=15)

    for _ in range(0, 3 * ceil(vol / 2), 3):
        table.cell(_, 0).merge(table.cell(_ + 1, 0))
        table.cell(_, 9).merge(table.cell(_ + 1, 9))
    per = 0
    for i in range(0, 3 * ceil(vol / 2), 3):
        for j in range(0, 13, 9):
            try:
                a1, b1, a2, b2, lambda1, lambda2 = mas_var[per]
            except IndexError:
                break
            per += 1
            table.rows[i].cells[j].paragraphs[0].add_run().add_picture('skoba.png', width=Cm(0.34), height=Cm(1.76))
            table.rows[i].cells[j + 1].text = "x`="
            table.rows[i].cells[j + 2].text = str(a1) + "x"
            if b1 > 0:
                table.rows[i].cells[j + 3].text = "+"
                table.rows[i].cells[j + 4].text = str(b1) + "y"
            elif b1 < 0:
                table.rows[i].cells[j + 3].text = "-"
                table.rows[i].cells[j + 4].text = str(abs(b1)) + "y"
            else:
                pass
            table.rows[i + 1].cells[j + 1].text = "y`="
            table.rows[i + 1].cells[j + 2].text = str(a2) + "x"
            if b2 > 0:
                table.rows[i + 1].cells[j + 3].text = "+"
                table.rows[i + 1].cells[j + 4].text = str(b2) + "y"
            elif b2 < 0:
                table.rows[i + 1].cells[j + 3].text = "-"
                table.rows[i + 1].cells[j + 4].text = str(abs(b2)) + "y"
            else:
                pass
    doc.save('../SDY_2_noANS.docx')


def write_to_file_with_ans3(mas_var, vol):
    doc = docx.Document()
    doc.sections[0].page_height = Cm(29.7)
    doc.sections[0].page_width = Cm(21.0)
    doc.sections[0].left_margin = Cm(1.5)
    doc.sections[0].right_margin = Cm(1.5)
    doc.sections[0].top_margin = Cm(1.5)
    doc.sections[0].bottom_margin = Cm(1.5)
    table = doc.add_table(rows=5 * ceil(vol / 2), cols=19)

    for _ in range(0, 5 * ceil(vol / 2), 5):
        table.cell(_, 0).merge(table.cell(_ + 2, 0))
        table.cell(_, 10).merge(table.cell(_ + 2, 10))
        table.cell(_ + 3, 1).merge(table.cell(_ + 3, 2))
        table.cell(_ + 3, 4).merge(table.cell(_ + 3, 5))
        table.cell(_ + 3, 7).merge(table.cell(_ + 3, 8))
        table.cell(_ + 3, 11).merge(table.cell(_ + 3, 12))
        table.cell(_ + 3, 14).merge(table.cell(_ + 3, 15))
        table.cell(_ + 3, 17).merge(table.cell(_ + 3, 18))
    per = 0
    for i in range(0, 5 * ceil(vol / 2), 5):
        for j in range(0, 18, 10):
            try:
                a1, b1, c1, a2, b2, c2, a3, b3, c3, lambda1, lambda2, lambda3 = mas_var[per]
            except IndexError:
                break
            per += 1
            table.rows[i].cells[j].paragraphs[0].add_run().add_picture('skoba.png', width=Cm(0.51), height=Cm(2.63))
            table.rows[i].cells[j + 1].text = "x`="
            table.rows[i].cells[j + 2].text = str(a1) + "x"
            if b1 > 0:
                table.rows[i].cells[j + 3].text = "+"
                table.rows[i].cells[j + 4].text = str(b1) + "y"
            elif b1 < 0:
                table.rows[i].cells[j + 3].text = "-"
                table.rows[i].cells[j + 4].text = str(abs(b1)) + "y"
            else:
                pass
            if c1 > 0:
                table.rows[i].cells[j + 5].text = "+"
                table.rows[i].cells[j + 6].text = str(c1) + "z"
            elif c1 < 0:
                table.rows[i].cells[j + 5].text = "-"
                table.rows[i].cells[j + 6].text = str(abs(c1)) + "z"
            else:
                pass

            table.rows[i + 1].cells[j + 1].text = "y`="
            table.rows[i + 1].cells[j + 2].text = str(a2) + "x"
            if b2 > 0:
                table.rows[i + 1].cells[j + 3].text = "+"
                table.rows[i + 1].cells[j + 4].text = str(b2) + "y"
            elif b2 < 0:
                table.rows[i + 1].cells[j + 3].text = "-"
                table.rows[i + 1].cells[j + 4].text = str(abs(b2)) + "y"
            else:
                pass
            if c2 > 0:
                table.rows[i + 1].cells[j + 5].text = "+"
                table.rows[i + 1].cells[j + 6].text = str(c2) + "z"
            elif c2 < 0:
                table.rows[i + 1].cells[j + 5].text = "-"
                table.rows[i + 1].cells[j + 6].text = str(abs(c2)) + "z"
            else:
                pass

            table.rows[i + 2].cells[j + 1].text = "z`="
            table.rows[i + 2].cells[j + 2].text = str(a3) + "x"
            if b3 > 0:
                table.rows[i + 2].cells[j + 3].text = "+"
                table.rows[i + 2].cells[j + 4].text = str(b3) + "y"
            elif b3 < 0:
                table.rows[i + 2].cells[j + 3].text = "-"
                table.rows[i + 2].cells[j + 4].text = str(abs(b3)) + "y"
            else:
                pass
            if c3 > 0:
                table.rows[i + 2].cells[j + 5].text = "+"
                table.rows[i + 2].cells[j + 6].text = str(c3) + "z"
            elif c3 < 0:
                table.rows[i + 2].cells[j + 5].text = "-"
                table.rows[i + 2].cells[j + 6].text = str(abs(c3)) + "z"
            else:
                pass

            table.rows[i + 3].cells[j].text = "\u03bb1="
            table.rows[i + 3].cells[j + 1].text = str(lambda1)
            table.rows[i + 3].cells[j + 3].text = "\u03bb2="
            table.rows[i + 3].cells[j + 4].text = str(lambda2)
            table.rows[i + 3].cells[j + 6].text = "\u03bb3="
            table.rows[i + 3].cells[j + 7].text = str(lambda3)
    doc.save('../SDY_3_ANS.docx')


def write_to_file_without_ans3(mas_var, vol):
    doc = docx.Document()
    doc.sections[0].page_height = Cm(29.7)
    doc.sections[0].page_width = Cm(21.0)
    doc.sections[0].left_margin = Cm(1.5)
    doc.sections[0].right_margin = Cm(1.5)
    doc.sections[0].top_margin = Cm(1.5)
    doc.sections[0].bottom_margin = Cm(1.5)
    table = doc.add_table(rows=4 * ceil(vol / 2), cols=19)

    for _ in range(0, 4 * ceil(vol / 2), 4):
        table.cell(_, 0).merge(table.cell(_ + 2, 0))
        table.cell(_, 10).merge(table.cell(_ + 2, 10))
    per = 0
    for i in range(0, 4 * ceil(vol / 2), 4):
        for j in range(0, 18, 10):
            try:
                a1, b1, c1, a2, b2, c2, a3, b3, c3, lambda1, lambda2, lambda3 = mas_var[per]
            except IndexError:
                break
            per += 1
            table.rows[i].cells[j].paragraphs[0].add_run().add_picture('skoba.png', width=Cm(0.51), height=Cm(2.63))
            table.rows[i].cells[j + 1].text = "x`="
            table.rows[i].cells[j + 2].text = str(a1) + "x"
            if b1 > 0:
                table.rows[i].cells[j + 3].text = "+"
                table.rows[i].cells[j + 4].text = str(b1) + "y"
            elif b1 < 0:
                table.rows[i].cells[j + 3].text = "-"
                table.rows[i].cells[j + 4].text = str(abs(b1)) + "y"
            else:
                pass
            if c1 > 0:
                table.rows[i].cells[j + 5].text = "+"
                table.rows[i].cells[j + 6].text = str(c1) + "z"
            elif c1 < 0:
                table.rows[i].cells[j + 5].text = "-"
                table.rows[i].cells[j + 6].text = str(abs(c1)) + "z"
            else:
                pass

            table.rows[i + 1].cells[j + 1].text = "y`="
            table.rows[i + 1].cells[j + 2].text = str(a2) + "x"
            if b2 > 0:
                table.rows[i + 1].cells[j + 3].text = "+"
                table.rows[i + 1].cells[j + 4].text = str(b2) + "y"
            elif b2 < 0:
                table.rows[i + 1].cells[j + 3].text = "-"
                table.rows[i + 1].cells[j + 4].text = str(abs(b2)) + "y"
            else:
                pass
            if c2 > 0:
                table.rows[i + 1].cells[j + 5].text = "+"
                table.rows[i + 1].cells[j + 6].text = str(c2) + "z"
            elif c2 < 0:
                table.rows[i + 1].cells[j + 5].text = "-"
                table.rows[i + 1].cells[j + 6].text = str(abs(c2)) + "z"
            else:
                pass

            table.rows[i + 2].cells[j + 1].text = "z`="
            table.rows[i + 2].cells[j + 2].text = str(a3) + "x"
            if b3 > 0:
                table.rows[i + 2].cells[j + 3].text = "+"
                table.rows[i + 2].cells[j + 4].text = str(b3) + "y"
            elif b3 < 0:
                table.rows[i + 2].cells[j + 3].text = "-"
                table.rows[i + 2].cells[j + 4].text = str(abs(b3)) + "y"
            else:
                pass
            if c3 > 0:
                table.rows[i + 2].cells[j + 5].text = "+"
                table.rows[i + 2].cells[j + 6].text = str(c3) + "z"
            elif c3 < 0:
                table.rows[i + 2].cells[j + 5].text = "-"
                table.rows[i + 2].cells[j + 6].text = str(abs(c3)) + "z"
            else:
                pass
    doc.save('../SDY_3_noANS.docx')
